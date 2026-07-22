import os
import re
import hashlib
from typing import Optional

def extract_text(file_path: str, filename: str = None) -> Optional[str]:
    """
    Extract text from any supported file format.
    Supports: PDF, DOCX, TXT, CSV, XLSX, XLS
    Returns clean text or None if extraction fails.
    """
    ext = (filename or file_path).lower().split(".")[-1]

    if ext == "pdf":
        return _extract_pdf(file_path)
    elif ext == "docx":
        return _extract_docx(file_path)
    elif ext == "txt":
        return _extract_txt(file_path)
    elif ext in ("xlsx", "xls"):
        return _extract_excel(file_path)
    elif ext == "csv":
        return _extract_csv(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

def _extract_pdf(file_path: str) -> str:
    import pymupdf4llm
    text = pymupdf4llm.to_markdown(file_path)
    return _clean_text(text)

def _extract_docx(file_path: str) -> str:
    from docx import Document
    doc = Document(file_path)
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            # Preserve heading structure
            if para.style.name.startswith("Heading"):
                parts.append(f"\n## {para.text.strip()}\n")
            else:
                parts.append(para.text.strip())

    # Extract tables
    for table in doc.tables:
        parts.append("\n")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
        parts.append("\n")

    return _clean_text("\n".join(parts))

def _extract_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return _clean_text(f.read())

def _extract_excel(file_path: str) -> str:
    import pandas as pd
    parts = []

    xl = pd.ExcelFile(file_path)
    for sheet_name in xl.sheet_names:
        df = pd.read_excel(file_path, sheet_name=sheet_name)
        if df.empty:
            continue

        parts.append(f"\n## {sheet_name}\n")

        # Convert to readable text
        for _, row in df.iterrows():
            row_text = " | ".join(
                f"{col}: {val}"
                for col, val in row.items()
                if str(val).strip() and str(val) != "nan"
            )
            if row_text.strip():
                parts.append(row_text)

    return _clean_text("\n".join(parts))

def _extract_csv(file_path: str) -> str:
    import pandas as pd
    df = pd.read_csv(file_path, encoding="utf-8", errors="ignore")
    parts = []
    for _, row in df.iterrows():
        row_text = " | ".join(
            f"{col}: {val}"
            for col, val in row.items()
            if str(val).strip() and str(val) != "nan"
        )
        if row_text.strip():
            parts.append(row_text)
    return _clean_text("\n".join(parts))

def _clean_text(text: str) -> str:
    # Fix broken numbers
    text = re.sub(r"(\d),\s*\n\s*(\d)", r"\1,\2", text)
    # Fix rupee symbol
    text = re.sub(r"₹\s+(\d)", r"₹\1", text)
    # Remove excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove page numbers
    text = re.sub(r"\n\s*\d+\s*\n", "\n", text)
    return text.strip()

SUPPORTED_FORMATS = ["pdf", "docx", "txt", "xlsx", "xls", "csv"]

def is_supported(filename: str) -> bool:
    ext = filename.lower().split(".")[-1]
    return ext in SUPPORTED_FORMATS
def clean_tables_with_gpt(text: str) -> str:
    """
    Find markdown tables in text and convert them to
    readable sentences using GPT-4o mini.
    """
    import re
    import os
    from openai import OpenAI

    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

    # Find markdown tables
    table_pattern = re.compile(
        r'(\|.+\|[\s\S]*?)(?=\n\n|\n#|\Z)',
        re.MULTILINE
    )

    tables = table_pattern.findall(text)
    if not tables:
        return text

    print(f"  Found {len(tables)} tables — cleaning with GPT-4o mini...")

    for table in tables:
        if len(table.strip()) < 50:
            continue

        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{
                    "role": "user",
                    "content": f"""Convert this table from an Indian tax document into clear, readable sentences. 
Each row should become one or two sentences.
Include all numbers, rates, thresholds and section numbers.
Do not add any information not in the table.
Do not use bullet points or lists — use plain sentences only.

Table:
{table}

Output as plain text sentences:"""
                }],
                temperature=0,
                max_tokens=1000
            )
            clean = response.choices[0].message.content.strip()
            text = text.replace(table, clean)
            print(f"  Cleaned table ({len(table)} → {len(clean)} chars)")
        except Exception as e:
            print(f"  Table cleaning failed: {e}")
            continue

    return text
